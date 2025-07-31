from pathlib import Path
from typing import List

import PyPDF2
import pandas as pd
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument

from config.settings import CHUNK_SIZE, CHUNK_OVERLAP, SUPPORTED_EXTENSIONS


class DocumentLoader:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            length_function=len,
        )

    def load_documents_from_directory(self, directory_path: str) -> List[LangchainDocument]:
        """Load and process documents from directory"""
        documents = []
        directory_path = Path(directory_path)

        if not directory_path.exists():
            return documents

        for file_path in directory_path.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_EXTENSIONS:
                content = self._load_file_content(file_path)
                if content:
                    metadata = {
                        "source": str(file_path),
                        "filename": file_path.name,
                        "directory": str(file_path.parent),
                        "extension": file_path.suffix.lower()
                    }

                    chunks = self.text_splitter.split_text(content)
                    for i, chunk in enumerate(chunks):
                        chunk_metadata = metadata.copy()
                        chunk_metadata["chunk_index"] = i
                        chunk_metadata["total_chunks"] = len(chunks)

                        documents.append(LangchainDocument(
                            page_content=chunk,
                            metadata=chunk_metadata
                        ))

        return documents

    def _load_file_content(self, file_path: Path) -> str:
        """Load content from file based on extension"""
        try:
            extension = file_path.suffix.lower()

            if extension in ['.txt', '.md']:
                return self._load_text_file(file_path)
            elif extension == '.pdf':
                return self._load_pdf_file(file_path)
            elif extension == '.docx':
                return self._load_docx_file(file_path)
            elif extension == '.csv':
                return self._load_csv_file(file_path)

            return ""
        except:
            return ""

    def _load_text_file(self, file_path: Path) -> str:
        """Load text files"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read()

    def _load_pdf_file(self, file_path: Path) -> str:
        """Load PDF files"""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text

    def _load_docx_file(self, file_path: Path) -> str:
        """Load DOCX files"""
        doc = Document(file_path)
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)

    def _load_csv_file(self, file_path: Path) -> str:
        """Load CSV files"""
        df = pd.read_csv(file_path)
        text = f"File: {file_path.name}\n"
        text += f"Columns: {', '.join(df.columns.tolist())}\n"
        text += f"Rows: {len(df)}\n\n"
        text += df.to_string(index=False)
        return text
